"""
从法律 docx 按条提取为 jsonl，并支持一键更新 references 数据库。

转换规则：
1. law_name: 段落 #2 的 text，去掉换行
2. 新条: runs[0] 以"第"开头且 runs[1] 以 \\u3000 开头
3. 章节标题（单 run 含 \\u3000）→ 跳过
4. 续段 → 用 \\n 拼到上一条

一键更新流程（相对路径以本脚本所在目录为基准）：
1. 对比 original_docx/ 与 references/ 文件名（A.docx ↔ A.jsonl）
2. A.docx 有对应 A.jsonl → 忽略
3. B.jsonl 没有对应 B.docx → 从 references/ 删除 B.jsonl
4. C.docx 没有对应 C.jsonl → 运行转换：original_docx/C.docx → references/C.jsonl
5. 全部完成后，生成 references/all_reference.jsonl（若已存在则先删再建）

用法：
    python docx_to_reference.py              # 一键更新
    python docx_to_reference.py <docx或目录>   # 仅转换（输出到 references/）
"""

import json
import os
import re
import sys
from docx import Document

# 相对路径：以本文件所在目录为基准
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ORIGINAL_DOCX_DIR = os.path.join(SCRIPT_DIR, "original_docx")
REFERENCES_DIR = os.path.join(SCRIPT_DIR, "references")
ALL_REFERENCE_FILE = "all_reference.jsonl"


def extract_law_name(doc, docx_path):
    """从段落 #2 提取 law_name，去掉换行"""
    if len(doc.paragraphs) > 2:
        name = doc.paragraphs[2].text.strip().replace("\n", "")
    else:
        name = os.path.basename(docx_path).rsplit(".", 1)[0]
        name = re.sub(r"_\d{8}$", "", name)
    return name


def is_new_article(para):
    runs = para.runs
    if len(runs) >= 2:
        r0 = runs[0].text.strip()
        r1 = runs[1].text
        if r0.startswith("第") and r1.startswith("\u3000"):
            return True
    return False


def is_chapter_heading(para):
    runs = para.runs
    if len(runs) == 1 and "\u3000" in runs[0].text:
        return True
    return False


def get_article_name(para):
    return para.runs[0].text.strip()


def get_article_text(para):
    text = "".join(run.text for run in para.runs[1:])
    return text.lstrip("\u3000").strip()


def extract_articles(docx_path):
    doc = Document(docx_path)
    law_name = extract_law_name(doc, docx_path)
    articles = []
    current_article = None
    current_text = None
    started = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or not para.runs:
            continue
        if is_new_article(para):
            if not started:
                started = True
            if current_article is not None:
                articles.append({"law_name": law_name, "article": current_article, "text": current_text})
            current_article = get_article_name(para)
            current_text = get_article_text(para)
            continue
        if not started or is_chapter_heading(para):
            continue
        if current_article is not None:
            current_text += "\n" + text

    if current_article is not None:
        articles.append({"law_name": law_name, "article": current_article, "text": current_text})
    return articles


def convert_file(docx_path, out_path=None):
    """转换单个 docx → jsonl，若未指定 out_path 则输出到 references/ 下同名 .jsonl"""
    articles = extract_articles(docx_path)
    if out_path is None:
        base = os.path.basename(docx_path).rsplit(".", 1)[0] + ".jsonl"
        out_path = os.path.join(REFERENCES_DIR, base)
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        for a in articles:
            f.write(json.dumps(a, ensure_ascii=False) + "\n")
    print(f"  转换: {os.path.basename(docx_path)} → {os.path.basename(out_path)}  ({len(articles)} 条)")
    return out_path


def get_base_name(path, ext):
    """文件名去掉扩展名，用于对比 A.docx / A.jsonl"""
    name = os.path.basename(path)
    if name.endswith(ext):
        return name[: -len(ext)]
    return None


def one_click_update():
    """一键更新：对比 original_docx 与 references，删多余 jsonl、补转缺的 docx，最后生成 all_reference.jsonl"""
    print("=" * 60)
    print("一键更新 references...", flush=True)
    if not os.path.isdir(ORIGINAL_DOCX_DIR):
        print(f"未找到目录: {ORIGINAL_DOCX_DIR}", flush=True)
        sys.exit(1)
    os.makedirs(REFERENCES_DIR, exist_ok=True)

    docx_files = [
        f for f in os.listdir(ORIGINAL_DOCX_DIR)
        if f.lower().endswith(".docx")
    ]
    jsonl_files = [
        f for f in os.listdir(REFERENCES_DIR)
        if f.endswith(".jsonl") and f != ALL_REFERENCE_FILE
    ]

    docx_bases = {get_base_name(f, ".docx"): f for f in docx_files}
    jsonl_bases = {get_base_name(f, ".jsonl"): f for f in jsonl_files}
    jsonl_bases_orig = dict(jsonl_bases)  # 保存副本，汇总时用

    # 1. 删除 references 里没有对应 docx 的 jsonl
    to_remove = [b for b in jsonl_bases if b not in docx_bases]
    for base in to_remove:
        path = os.path.join(REFERENCES_DIR, jsonl_bases[base])
        os.remove(path)
        print(f"  删除（无对应 docx）: {jsonl_bases[base]}")
    for base in to_remove:
        del jsonl_bases[base]

    # 2. 对没有 jsonl 的 docx 做转换
    to_convert = [b for b in docx_bases if b not in jsonl_bases]
    if to_convert:
        print(f"  待转换 {len(to_convert)} 个 docx:")
        for base in sorted(to_convert):
            docx_path = os.path.join(ORIGINAL_DOCX_DIR, docx_bases[base])
            out_path = os.path.join(REFERENCES_DIR, base + ".jsonl")
            convert_file(docx_path, out_path)

    # 3. 生成 all_reference.jsonl
    all_path = os.path.join(REFERENCES_DIR, ALL_REFERENCE_FILE)
    if os.path.isfile(all_path):
        os.remove(all_path)
        print(f"  已删除旧 {ALL_REFERENCE_FILE}")

    jsonl_list = [
        f for f in os.listdir(REFERENCES_DIR)
        if f.endswith(".jsonl") and f != ALL_REFERENCE_FILE
    ]
    with open(all_path, "w", encoding="utf-8") as out:
        for name in sorted(jsonl_list):
            path = os.path.join(REFERENCES_DIR, name)
            with open(path, "r", encoding="utf-8") as inp:
                for line in inp:
                    line = line.strip()
                    if line:
                        out.write(line + "\n")

    # 4. 汇总报告
    print("本次更新汇总:")
    print("=" * 60)
    if to_remove:
        print(f"\n  删除了 {len(to_remove)} 个 jsonl（无对应 docx）:")
        for base in sorted(to_remove):
            print(f"    - {jsonl_bases_orig[base]}")
    else:
        print("\n  删除: 无")
    if to_convert:
        print(f"\n  新增了 {len(to_convert)} 个 jsonl（从 docx 转换）:")
        for base in sorted(to_convert):
            print(f"    + {base}.jsonl")
    else:
        print("\n  新增: 无")
    unchanged = len(docx_bases) - len(to_convert)
    print(f"\n  未变动: {unchanged} 个")
    print(f"  references/ 下当前共 {len(jsonl_list)} 个文档")
    print("=" * 60)


def main():
    if len(sys.argv) == 1:
        one_click_update()
        return

    target = sys.argv[1]
    if not os.path.exists(target):
        print(f"找不到: {target}")
        sys.exit(1)

    os.makedirs(REFERENCES_DIR, exist_ok=True)
    if os.path.isfile(target) and target.lower().endswith(".docx"):
        convert_file(target)
    elif os.path.isdir(target):
        files = [os.path.join(target, f) for f in sorted(os.listdir(target)) if f.lower().endswith(".docx")]
        print(f"批量转换 {len(files)} 个 docx 到 {REFERENCES_DIR}:")
        for f in files:
            convert_file(f)
    else:
        print("请传入 .docx 文件或包含 .docx 的目录")
        sys.exit(1)


if __name__ == "__main__":
    main()
